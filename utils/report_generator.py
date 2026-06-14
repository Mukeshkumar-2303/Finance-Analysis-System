import os
import re
import datetime
from xml.sax.saxutils import escape

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    Image,
    PageBreak
)


# =========================================================
# COMMON HELPERS
# =========================================================

def clean_text(text):
    text = str(text)
    text = text.replace("₹", "INR ")
    text = text.replace("■", "INR ")
    text = text.replace("**", "")
    text = text.replace("#", "")
    text = text.replace("–", "-")
    text = text.replace("—", "-")
    return text.strip()


def is_markdown_table_line(line):
    line = line.strip()
    return line.startswith("|") and line.endswith("|") and line.count("|") >= 2


def is_markdown_separator_line(line):
    clean = line.replace("|", "").replace(":", "").replace("-", "").strip()
    return clean == ""


def parse_markdown_table(table_lines):
    rows = []

    for line in table_lines:
        line = line.strip()

        if not is_markdown_table_line(line):
            continue

        if is_markdown_separator_line(line):
            continue

        cells = [cell.strip() for cell in line.strip("|").split("|")]
        rows.append(cells)

    if len(rows) < 2:
        return []

    return rows


# =========================================================
# DOCX REPORT HELPERS
# =========================================================

def add_dataframe_table_to_docx(doc, title, dataframe):
    doc.add_heading(title, level=2)

    if dataframe.empty:
        doc.add_paragraph("No data available.")
        return

    table = doc.add_table(rows=1, cols=len(dataframe.columns))
    table.style = "Table Grid"

    for i, col in enumerate(dataframe.columns):
        table.rows[0].cells[i].text = str(col)

    for _, row in dataframe.iterrows():
        cells = table.add_row().cells

        for i, col in enumerate(dataframe.columns):
            cells[i].text = str(row[col])

    doc.add_paragraph("")


def add_markdown_table_to_docx(doc, table_rows):
    if not table_rows:
        return

    col_count = max(len(row) for row in table_rows)

    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"

    header_cells = table.rows[0].cells

    for i in range(col_count):
        header_cells[i].text = table_rows[0][i] if i < len(table_rows[0]) else ""

    for row in table_rows[1:]:
        cells = table.add_row().cells

        for i in range(col_count):
            cells[i].text = row[i] if i < len(row) else ""

    doc.add_paragraph("")


def add_ai_report_to_docx(doc, ai_report):
    lines = ai_report.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Convert markdown tables into real Word tables
        if is_markdown_table_line(line):
            table_lines = []

            while i < len(lines) and is_markdown_table_line(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1

            table_rows = parse_markdown_table(table_lines)
            add_markdown_table_to_docx(doc, table_rows)
            continue

        clean_line = clean_text(line)

        if re.match(r"^\d+\. ", clean_line):
            doc.add_heading(clean_line, level=2)
        elif re.match(r"^\d+\.\d+ ", clean_line):
            doc.add_heading(clean_line, level=3)
        elif clean_line.startswith("* "):
            doc.add_paragraph(clean_line.replace("* ", "", 1), style="List Bullet")
        elif clean_line.startswith("- "):
            doc.add_paragraph(clean_line.replace("- ", "", 1), style="List Bullet")
        else:
            doc.add_paragraph(clean_line)

        i += 1


def create_docx_report(
    output_path,
    structured_data,
    metrics_df,
    ratios_df,
    chart_paths,
    ai_report,
    liquidity_rating,
    liquidity_rating_reason
):
    doc = Document()

    title = doc.add_heading("FIN AI Financial Statement Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Generated on: {datetime.datetime.now().strftime('%d-%m-%Y %H:%M')}")
    doc.add_paragraph(
        "Prepared using Python, AI, document extraction, table extraction, "
        "financial ratio analysis, and infographic generation."
    )

    doc.add_heading("1. Document Classification", level=1)

    classification = structured_data.get("classification", {})

    for key, value in classification.items():
        p = doc.add_paragraph()
        p.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
        p.add_run(str(value))

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Liquidity Risk Rating: ").bold = True
    p.add_run(liquidity_rating)

    p = doc.add_paragraph()
    p.add_run("Liquidity Rating Reason: ").bold = True
    p.add_run(liquidity_rating_reason)

    add_dataframe_table_to_docx(doc, "2. Extracted Financial Metrics", metrics_df)
    add_dataframe_table_to_docx(doc, "3. Calculated Financial Ratios", ratios_df)

    doc.add_heading("4. Infographic Charts", level=1)

    for chart_path in chart_paths:
        if os.path.exists(chart_path):
            doc.add_picture(chart_path, width=Inches(6))
            doc.add_paragraph(os.path.basename(chart_path))

    doc.add_heading("5. AI Generated Financial Analysis", level=1)

    add_ai_report_to_docx(doc, ai_report)

    doc.add_heading("6. Disclaimer", level=1)
    doc.add_paragraph(
        "This report is generated automatically from uploaded financial documents. "
        "The analysis depends on the quality and completeness of extracted text and tables. "
        "All figures should be reviewed before business, board, credit, or investment decisions."
    )

    doc.save(output_path)

    return output_path


# =========================================================
# PDF REPORT HELPERS
# =========================================================

def pdf_para(text, style):
    return Paragraph(escape(clean_text(text)), style)


def create_pdf_table(title, dataframe, styles):
    elements = []

    elements.append(pdf_para(title, styles["Heading2"]))
    elements.append(Spacer(1, 8))

    if dataframe.empty:
        elements.append(pdf_para("No data available.", styles["Normal"]))
        return elements

    df = dataframe.copy()

    if "Source Quote" in df.columns:
        df["Source Quote"] = df["Source Quote"].astype(str).apply(lambda x: x[:90])

    table_data = []

    header = [pdf_para(col, styles["Normal"]) for col in df.columns]
    table_data.append(header)

    for _, row in df.iterrows():
        table_data.append([
            pdf_para(str(row[col]), styles["Normal"])
            for col in df.columns
        ])

    page_width = A4[0] - 60
    col_count = len(df.columns)
    col_width = page_width / col_count

    table = Table(table_data, colWidths=[col_width] * col_count)

    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTSIZE", (0, 0), (-1, -1), 6),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
    ]))

    elements.append(table)
    elements.append(Spacer(1, 14))

    return elements


def create_pdf_markdown_table(table_rows, styles):
    elements = []

    if not table_rows:
        return elements

    col_count = max(len(row) for row in table_rows)
    page_width = A4[0] - 60

    # Give first column more space
    if col_count == 4:
        col_widths = [
            page_width * 0.34,
            page_width * 0.22,
            page_width * 0.22,
            page_width * 0.22,
        ]
    elif col_count == 3:
        col_widths = [
            page_width * 0.40,
            page_width * 0.30,
            page_width * 0.30,
        ]
    else:
        col_widths = [page_width / col_count] * col_count

    table_data = []

    for row in table_rows:
        normalized_row = row + [""] * (col_count - len(row))
        table_data.append([
            pdf_para(cell, styles["Normal"])
            for cell in normalized_row
        ])

    table = Table(table_data, colWidths=col_widths, repeatRows=1)

    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbeafe")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#94a3b8")),
        ("BACKGROUND", (0, 1), (-1, -1), colors.white),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTSIZE", (0, 0), (-1, -1), 7),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))

    elements.append(table)
    elements.append(Spacer(1, 12))

    return elements


def add_ai_report_to_pdf(elements, ai_report, styles):
    lines = ai_report.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            i += 1
            continue

        # Convert markdown tables into real PDF tables
        if is_markdown_table_line(line):
            table_lines = []

            while i < len(lines) and is_markdown_table_line(lines[i].strip()):
                table_lines.append(lines[i].strip())
                i += 1

            table_rows = parse_markdown_table(table_lines)
            elements.extend(create_pdf_markdown_table(table_rows, styles))
            continue

        clean_line = clean_text(line)

        if re.match(r"^\d+\. ", clean_line):
            elements.append(pdf_para(clean_line, styles["Heading2"]))
        elif re.match(r"^\d+\.\d+ ", clean_line):
            elements.append(pdf_para(clean_line, styles["Heading3"]))
        elif clean_line.startswith("* "):
            elements.append(pdf_para("• " + clean_line.replace("* ", "", 1), styles["Normal"]))
        elif clean_line.startswith("- "):
            elements.append(pdf_para("• " + clean_line.replace("- ", "", 1), styles["Normal"]))
        else:
            elements.append(pdf_para(clean_line, styles["Normal"]))

        elements.append(Spacer(1, 4))
        i += 1


def create_pdf_report(
    output_path,
    structured_data,
    metrics_df,
    ratios_df,
    chart_paths,
    ai_report,
    liquidity_rating,
    liquidity_rating_reason
):
    pdf = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=30,
        leftMargin=30,
        topMargin=30,
        bottomMargin=30
    )

    styles = getSampleStyleSheet()
    elements = []

    elements.append(pdf_para("FIN AI Financial Statement Analysis Report", styles["Title"]))
    elements.append(Spacer(1, 12))

    elements.append(pdf_para(f"Generated on: {datetime.datetime.now().strftime('%d-%m-%Y %H:%M')}", styles["Normal"]))
    elements.append(pdf_para(
        "Prepared using Python, AI, document extraction, table extraction, ratio analysis, and infographic generation.",
        styles["Normal"]
    ))
    elements.append(Spacer(1, 12))

    elements.append(pdf_para("1. Document Classification", styles["Heading1"]))

    classification = structured_data.get("classification", {})

    for key, value in classification.items():
        elements.append(pdf_para(f"{key.replace('_', ' ').title()}: {value}", styles["Normal"]))

    elements.append(Spacer(1, 8))
    elements.append(pdf_para(f"Liquidity Risk Rating: {liquidity_rating}", styles["Normal"]))
    elements.append(pdf_para(f"Liquidity Rating Reason: {liquidity_rating_reason}", styles["Normal"]))
    elements.append(Spacer(1, 12))

    metrics_pdf_df = metrics_df[["Metric", "Current", "Previous", "Unit", "Confidence"]]
    ratios_pdf_df = ratios_df

    elements += create_pdf_table("2. Extracted Financial Metrics", metrics_pdf_df, styles)
    elements += create_pdf_table("3. Calculated Financial Ratios", ratios_pdf_df, styles)

    elements.append(PageBreak())

    elements.append(pdf_para("4. Infographic Charts", styles["Heading1"]))
    elements.append(Spacer(1, 12))

    for chart_path in chart_paths:
        if os.path.exists(chart_path):
            elements.append(Image(chart_path, width=6.4 * inch, height=3.2 * inch))
            elements.append(Spacer(1, 12))

    elements.append(PageBreak())

    elements.append(pdf_para("5. AI Generated Financial Analysis", styles["Heading1"]))
    elements.append(Spacer(1, 10))

    add_ai_report_to_pdf(elements, ai_report, styles)

    elements.append(PageBreak())

    elements.append(pdf_para("6. Disclaimer", styles["Heading1"]))
    elements.append(pdf_para(
        "This report is generated automatically from uploaded financial documents. "
        "The analysis depends on the quality and completeness of extracted text and tables. "
        "All figures should be reviewed before business, board, credit, or investment decisions.",
        styles["Normal"]
    ))

    pdf.build(elements)

    return output_path