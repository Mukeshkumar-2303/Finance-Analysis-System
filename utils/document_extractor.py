# document_extractor.py
import os
import fitz
import pdfplumber
from docx import Document


def extract_text_from_pdf(file_path: str):
    text = ""
    page_texts = {}

    try:
        pdf = fitz.open(file_path)

        for page_number, page in enumerate(pdf, start=1):
            page_text = page.get_text("text")
            page_texts[page_number] = page_text

            text += f"\n\n--- PAGE {page_number} ---\n"
            text += page_text

        pdf.close()

    except Exception as e:
        text += f"\n[ERROR READING PDF TEXT: {e}]"

    return text, page_texts


def extract_tables_with_pymupdf(file_path: str):
    table_text = ""

    try:
        pdf = fitz.open(file_path)

        for page_number, page in enumerate(pdf, start=1):
            try:
                tables = page.find_tables()

                if tables and tables.tables:
                    table_text += f"\n\n--- TABLES FROM PAGE {page_number} USING PYMUPDF ---\n"

                    for table_index, table in enumerate(tables.tables, start=1):
                        data = table.extract()
                        table_text += f"\nTABLE {table_index}\n"

                        for row in data:
                            clean_row = [
                                str(cell).strip().replace("\n", " ")
                                if cell is not None else ""
                                for cell in row
                            ]
                            table_text += " | ".join(clean_row) + "\n"

            except Exception:
                pass

        pdf.close()

    except Exception as e:
        table_text += f"\n[ERROR EXTRACTING TABLES WITH PYMUPDF: {e}]"

    return table_text


def extract_tables_with_pdfplumber(file_path: str):
    table_text = ""

    try:
        with pdfplumber.open(file_path) as pdf:
            for page_number, page in enumerate(pdf.pages, start=1):
                try:
                    tables = page.extract_tables()

                    if tables:
                        table_text += f"\n\n--- TABLES FROM PAGE {page_number} USING PDFPLUMBER ---\n"

                        for table_index, table in enumerate(tables, start=1):
                            table_text += f"\nTABLE {table_index}\n"

                            for row in table:
                                clean_row = [
                                    str(cell).strip().replace("\n", " ")
                                    if cell is not None else ""
                                    for cell in row
                                ]
                                table_text += " | ".join(clean_row) + "\n"

                except Exception:
                    pass

    except Exception as e:
        table_text += f"\n[ERROR EXTRACTING TABLES WITH PDFPLUMBER: {e}]"

    return table_text


def extract_text_and_tables_from_docx(file_path: str):
    text = ""
    table_text = ""

    try:
        doc = Document(file_path)

        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text.strip() + "\n"

        for table_index, table in enumerate(doc.tables, start=1):
            table_text += f"\n\n--- DOCX TABLE {table_index} ---\n"

            for row in table.rows:
                clean_row = []

                for cell in row.cells:
                    clean_row.append(cell.text.strip().replace("\n", " "))

                table_text += " | ".join(clean_row) + "\n"

    except Exception as e:
        text += f"\n[ERROR READING DOCX: {e}]"

    return text, table_text


def extract_text_from_txt(file_path: str):
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except Exception:
        try:
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()
        except Exception as e:
            return f"\n[ERROR READING TXT: {e}]"


def extract_document(file_path: str):
    extension = file_path.lower().split(".")[-1]

    if extension == "pdf":
        text, page_texts = extract_text_from_pdf(file_path)
        table_text_1 = extract_tables_with_pymupdf(file_path)
        table_text_2 = extract_tables_with_pdfplumber(file_path)
        table_text = table_text_1 + "\n\n" + table_text_2
        return text, table_text, page_texts

    if extension == "docx":
        text, table_text = extract_text_and_tables_from_docx(file_path)
        return text, table_text, {}

    if extension == "txt":
        text = extract_text_from_txt(file_path)
        return text, "", {}

    return f"[UNSUPPORTED FILE TYPE: {file_path}]", "", {}


IMPORTANT_KEYWORDS = [
    "balance sheet",
    "statement of profit and loss",
    "statement of cash flows",
    "cash flow statement",
    "current assets",
    "total current assets",
    "current liabilities",
    "total current liabilities",
    "cash and cash equivalents",
    "inventories",
    "trade receivables",
    "trade payables",
    "borrowings",
    "total assets",
    "total liabilities",
    "profit before tax",
    "profit after tax",
    "profit for the year",
    "net profit",
    "pat",
    "pbt",
    "revenue from operations",
    "gross revenue",
    "total income",
    "ebitda",
    "operating activities",
    "net cash generated from operating activities",
    "capital expenditure",
    "capex",
    "contingent liabilities",
]


def select_relevant_lines(text: str, keywords=None, window=4, max_chars=90000):
    if keywords is None:
        keywords = IMPORTANT_KEYWORDS

    lines = text.splitlines()
    selected = []

    for i, line in enumerate(lines):
        lower_line = line.lower()

        if any(keyword in lower_line for keyword in keywords):
            start = max(0, i - window)
            end = min(len(lines), i + window + 1)

            selected.append("\n--- MATCH CONTEXT ---")

            for j in range(start, end):
                selected.append(lines[j])

    result = "\n".join(selected)

    if len(result) < 10000:
        result = text[:max_chars]

    return result[:max_chars]


def select_relevant_table_rows(table_text: str, keywords=None, max_chars=90000):
    if keywords is None:
        keywords = IMPORTANT_KEYWORDS

    lines = table_text.splitlines()
    selected = []

    for line in lines:
        lower_line = line.lower()

        if any(keyword in lower_line for keyword in keywords):
            selected.append(line)

    result = "\n".join(selected)

    if len(result) < 5000:
        result = table_text[:max_chars]

    return result[:max_chars]